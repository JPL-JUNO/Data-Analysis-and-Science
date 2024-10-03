"""
@File         : a.py
@Author(s)    : Stephen CUI
@LastEditor(s): Stephen CUI
@CreatedTime  : 2024-09-26 22:53:55
@Email        : cuixuanstephen@gmail.com
@Description  : Python 实战 Word 案例
"""

import docx

my_doc = docx.Document("./DATA/prose.docx")
my_text = "再让我们回顾一下故事的开头，不正是乔的新邻居告诉他地里出黄金的吗?然而，事实上是乔对英国的语言理解得还不够透彻。他的新邻居其实是说他那块土地有肥沃的土壤，所以你应该知道黄金的概念来自哪儿了吧。"

my_doc.add_paragraph(my_text)
my_doc.save("./RES/prose.docx")

my_doc = docx.Document("../DATA/insert_paragraph_before.docx")
my_para2 = my_doc.paragraphs[1].insert_paragraph_before(
    "    Did Joe's sons become slaves to the digging? No, they were inspired because they had visions of what money could do for them and did not focus on the money itself."
)

my_para4 = my_doc.add_paragraph(
    "    乔的儿子们成为挖掘土地的奴隶了吗?不，他们只是一昧地幻想着金钱能为他们做什么，并被无数个幻想所激励，但从未考虑过获取金钱的正确途径。"
)
my_doc.save("../RES/insert_paragraph_before.docx")

my_doc = docx.Document("../DATA/remove.docx")
my_para1 = my_doc.paragraphs[1]._element
my_para1.getparent().remove(my_para1)

my_para2 = my_doc.paragraphs[2]._element
my_para2.getparent().remove(my_para2)
my_doc.save("remove.docx")
